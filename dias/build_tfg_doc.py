from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from textwrap import dedent
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = Path.home() / "Downloads"
OUTPUT_DOCX = OUTPUT_DIR / "TFG_PlotSkip_primera_version.docx"
ASSETS_DIR = ROOT / "dias" / "generated_doc_assets"
LOGO_PATH = ROOT / "Plotea.png"
FONT_REGULAR = r"C:\Windows\Fonts\arial.ttf"
FONT_BOLD = r"C:\Windows\Fonts\arialbd.ttf"

COLOR_DARK = "#0F172A"
COLOR_ACCENT = "#1F6FEB"
COLOR_LIGHT = "#EAF2FF"
COLOR_MID = "#475569"
COLOR_BORDER = "#B8C6DB"
COLOR_SUCCESS = "#0F766E"


@dataclass
class TableRow:
    values: list[str]


@dataclass
class UseCase:
    code: str
    title: str
    actor: str
    stakeholders: str
    description: str
    preconditions: str
    postconditions: str
    steps: list[str]
    alternatives: str
    frequency: str


def ensure_dirs() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size=size)


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: str, outline: str, text_fill: str = "#FFFFFF", radius: int = 16, bold: bool = False) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    font = load_font(26 if bold else 24, bold=bold)
    left, top, right, bottom = xy
    max_width = (right - left) - 30
    lines: list[str] = []
    current = ""
    for word in text.split():
        candidate = f"{current} {word}".strip()
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_height = 32
    total_height = line_height * len(lines)
    y = top + ((bottom - top - total_height) // 2)
    for line in lines:
        w = draw.textlength(line, font=font)
        x = left + ((right - left - w) // 2)
        draw.text((x, y), line, fill=text_fill, font=font)
        y += line_height


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str) -> None:
    draw.line([start, end], fill=fill, width=5)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        arrow = [(ex, ey), (ex - 16 * direction, ey - 9), (ex - 16 * direction, ey + 9)]
    else:
        direction = 1 if ey > sy else -1
        arrow = [(ex, ey), (ex - 9, ey - 16 * direction), (ex + 9, ey - 16 * direction)]
    draw.polygon(arrow, fill=fill)


def make_context_diagram() -> Path:
    path = ASSETS_DIR / "context_diagram.png"
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)

    title_font = load_font(34, bold=True)
    draw.text((60, 40), "Diagrama de contexto de PlotSkip", fill=COLOR_DARK, font=title_font)

    draw_box(draw, (590, 290, 1010, 560), "PlotSkip\nAplicacion movil + API", COLOR_ACCENT, COLOR_DARK, bold=True)
    draw_box(draw, (120, 160, 420, 310), "Usuario registrado", COLOR_SUCCESS, COLOR_DARK)
    draw_box(draw, (120, 560, 420, 710), "TMDB API", "#7C3AED", COLOR_DARK)
    draw_box(draw, (1170, 180, 1470, 330), "Servicio de correo\nResend", "#EA580C", COLOR_DARK)
    draw_box(draw, (1170, 390, 1470, 540), "PostgreSQL\nNeon", "#2563EB", COLOR_DARK)
    draw_box(draw, (1170, 600, 1470, 750), "Almacen externo\nde avatares\n(evolucion)", "#64748B", COLOR_DARK)

    draw_arrow(draw, (420, 235), (590, 235), COLOR_DARK)
    draw_arrow(draw, (590, 275), (420, 275), COLOR_DARK)
    draw_arrow(draw, (420, 635), (590, 635), COLOR_DARK)
    draw_arrow(draw, (1010, 255), (1170, 255), COLOR_DARK)
    draw_arrow(draw, (1010, 465), (1170, 465), COLOR_DARK)
    draw_arrow(draw, (1010, 675), (1170, 675), COLOR_DARK)

    note_font = load_font(22)
    draw.text((70, 800), "El sistema se apoya en TMDB para catalogo y metadatos, en un servicio de email para recuperacion de contrasena y en PostgreSQL para el almacenamiento persistente.", fill=COLOR_MID, font=note_font)

    image.save(path)
    return path


def make_use_case_diagram() -> Path:
    path = ASSETS_DIR / "use_case_diagram.png"
    image = Image.new("RGB", (1700, 1050), "white")
    draw = ImageDraw.Draw(image)

    draw.text((60, 36), "Modelo resumido de casos de uso", fill=COLOR_DARK, font=load_font(34, bold=True))
    actor_box = (70, 360, 330, 660)
    draw_box(draw, actor_box, "Usuario de PlotSkip", COLOR_SUCCESS, COLOR_DARK)

    system_box = (420, 110, 1590, 920)
    draw.rounded_rectangle(system_box, radius=24, outline=COLOR_DARK, width=4)
    draw.text((470, 140), "Sistema PlotSkip", fill=COLOR_DARK, font=load_font(28, bold=True))

    cases = [
        ((520, 220, 960, 300), "Registrarse e iniciar sesion"),
        ((1010, 220, 1450, 300), "Recuperar contrasena"),
        ((520, 350, 960, 430), "Buscar peliculas y series"),
        ((1010, 350, 1450, 430), "Marcar vista o watchlist"),
        ((520, 480, 960, 560), "Escribir y votar resenas"),
        ((1010, 480, 1450, 560), "Seguir usuarios y ver feed"),
        ((520, 610, 960, 690), "Crear listas y compartirlas"),
        ((1010, 610, 1450, 690), "Consultar perfil y estadisticas"),
        ((765, 740, 1205, 820), "Gestionar colaboracion en listas"),
    ]
    for xy, label in cases:
        draw.rounded_rectangle(xy, radius=35, outline=COLOR_ACCENT, width=4, fill=COLOR_LIGHT)
        font = load_font(22, bold=True)
        left, top, right, bottom = xy
        width = right - left - 24
        words = label.split()
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if draw.textlength(candidate, font=font) <= width:
                current = candidate
            else:
                lines.append(current)
                current = word
        if current:
            lines.append(current)
        y = top + ((bottom - top - (len(lines) * 28)) // 2)
        for line in lines:
            w = draw.textlength(line, font=font)
            x = left + ((right - left - w) // 2)
            draw.text((x, y), line, fill=COLOR_DARK, font=font)
            y += 28

    actor_points = [(330, 510), (330, 510), (330, 510), (330, 510), (330, 510), (330, 510), (330, 510), (330, 510), (330, 510)]
    targets = [(520, 260), (1010, 260), (520, 390), (1010, 390), (520, 520), (1010, 520), (520, 650), (1010, 650), (765, 780)]
    for start, end in zip(actor_points, targets):
        draw_arrow(draw, start, end, "#64748B")

    image.save(path)
    return path


def make_architecture_diagram() -> Path:
    path = ASSETS_DIR / "architecture_diagram.png"
    image = Image.new("RGB", (1700, 1000), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 36), "Arquitectura general del sistema", fill=COLOR_DARK, font=load_font(34, bold=True))

    draw_box(draw, (120, 350, 420, 560), "Cliente movil\nReact Native + Expo", COLOR_SUCCESS, COLOR_DARK, bold=True)
    draw_box(draw, (560, 320, 1120, 590), "Backend API\nFastAPI + Clean Architecture\nJWT + refresh tokens", COLOR_ACCENT, COLOR_DARK, bold=True)
    draw_box(draw, (1260, 170, 1560, 320), "TMDB API", "#7C3AED", COLOR_DARK)
    draw_box(draw, (1260, 390, 1560, 540), "PostgreSQL", "#2563EB", COLOR_DARK)
    draw_box(draw, (1260, 610, 1560, 760), "Resend / email", "#EA580C", COLOR_DARK)
    draw_box(draw, (620, 720, 1060, 850), "Capa de datos local segura\nSecureStore + AsyncStorage", "#334155", COLOR_DARK)

    draw_arrow(draw, (420, 455), (560, 455), COLOR_DARK)
    draw_arrow(draw, (560, 490), (420, 490), COLOR_DARK)
    draw_arrow(draw, (1120, 245), (1260, 245), COLOR_DARK)
    draw_arrow(draw, (1120, 465), (1260, 465), COLOR_DARK)
    draw_arrow(draw, (1120, 685), (1260, 685), COLOR_DARK)
    draw_arrow(draw, (760, 590), (760, 720), COLOR_DARK)

    image.save(path)
    return path


def make_analysis_class_diagram() -> Path:
    path = ASSETS_DIR / "analysis_class_diagram.png"
    image = Image.new("RGB", (1750, 1100), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 36), "Modelo de clases de analisis", fill=COLOR_DARK, font=load_font(34, bold=True))

    boxes = {
        "Usuario": (90, 170, 420, 430),
        "Resena": (520, 120, 890, 360),
        "WatchLog": (520, 430, 890, 670),
        "Lista": (1000, 120, 1350, 360),
        "ElementoLista": (1000, 430, 1450, 700),
        "Comentario": (90, 520, 420, 760),
        "MediaRef": (520, 760, 930, 980),
        "Follow": (1000, 790, 1350, 980),
    }
    text_map = {
        "Usuario": "Usuario\n- id\n- email\n- username\n- display_name\n- bio\n- avatar_url",
        "Resena": "Resena\n- id\n- rating\n- title\n- body\n- contains_spoilers\n- created_at",
        "WatchLog": "WatchLog\n- id\n- watched_at\n- rating\n- notes",
        "Lista": "Lista\n- id\n- name\n- description\n- is_public",
        "ElementoLista": "ElementoLista\n- position\n- added_at\n- added_by_user_id",
        "Comentario": "Comentario\n- id\n- body\n- created_at",
        "MediaRef": "MediaRef\n- tmdb_id\n- media_type\n- title\n- poster_path\n- release_year",
        "Follow": "Follow\n- follower_id\n- followed_id\n- created_at",
    }
    for name, xy in boxes.items():
        draw_box(draw, xy, text_map[name], COLOR_LIGHT, COLOR_ACCENT, text_fill=COLOR_DARK, bold=True)

    connections = [
        ((420, 240), (520, 240)),
        ((420, 610), (520, 610)),
        ((890, 240), (1000, 240)),
        ((890, 550), (1000, 550)),
        ((705, 360), (705, 430)),
        ((700, 670), (700, 760)),
        ((260, 430), (260, 520)),
        ((1180, 360), (1180, 430)),
        ((1180, 700), (1180, 790)),
    ]
    for start, end in connections:
        draw_arrow(draw, start, end, COLOR_DARK)

    image.save(path)
    return path


def make_design_class_diagram() -> Path:
    path = ASSETS_DIR / "design_class_diagram.png"
    image = Image.new("RGB", (1700, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 36), "Modelo de clases de diseno por capas", fill=COLOR_DARK, font=load_font(34, bold=True))

    layers = [
        ((110, 180, 1590, 300), "Presentacion", "Screens Expo Router, ViewModels, componentes compartidos"),
        ((110, 340, 1590, 460), "Dominio", "Entidades, contratos de repositorio y casos de uso"),
        ((110, 500, 1590, 620), "Datos", "Repositorios HTTP/SQLAlchemy, mapeadores, clientes externos"),
        ((110, 660, 1590, 780), "Infraestructura", "Config, base de datos, auth JWT, almacenamiento seguro, email"),
    ]
    fills = [COLOR_LIGHT, "#EEF6F0", "#FFF4E6", "#EDE9FE"]
    for (xy, title, body), fill in zip(layers, fills):
        draw.rounded_rectangle(xy, radius=20, fill=fill, outline=COLOR_DARK, width=3)
        draw.text((xy[0] + 24, xy[1] + 18), title, fill=COLOR_DARK, font=load_font(28, bold=True))
        draw.text((xy[0] + 24, xy[1] + 60), body, fill=COLOR_MID, font=load_font(22))

    draw_arrow(draw, (850, 300), (850, 340), COLOR_DARK)
    draw_arrow(draw, (850, 460), (850, 500), COLOR_DARK)
    draw_arrow(draw, (850, 620), (850, 660), COLOR_DARK)
    image.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shd)


def set_document_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Title", 24, COLOR_DARK),
        ("Heading 1", 16, COLOR_DARK),
        ("Heading 2", 13, COLOR_DARK),
        ("Heading 3", 11.5, COLOR_ACCENT),
        ("Heading 4", 10.8, COLOR_DARK),
        ("Heading 5", 10.5, COLOR_MID),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*hex_to_rgb(color))

    if "CaptionCustom" not in styles:
        cap = styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Arial"
        cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = RGBColor(*hex_to_rgb(COLOR_MID))
        cap.paragraph_format.space_after = Pt(8)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_paragraph.add_run("Juan Caro Vaquero, Proyecto DAM, Curso 2025-2026")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(*hex_to_rgb(COLOR_MID))

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run("- ")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    add_page_field(footer_paragraph)
    run2 = footer_paragraph.add_run(" -")
    run2.font.name = "Arial"
    run2.font.size = Pt(8.5)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Actualiza el indice en Word o LibreOffice si fuera necesario."
    fld_separate.append(txt)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def add_paragraphs(document: Document, text: str) -> None:
    for paragraph_text in [part.strip() for part in dedent(text).strip().split("\n\n") if part.strip()]:
        p = document.add_paragraph(paragraph_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_numbered(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[TableRow], widths_cm: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[idx], COLOR_ACCENT)
        for paragraph in hdr_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = "Arial"
                run.font.size = Pt(9)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row.values):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[idx].paragraphs:
                if idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)

    if widths_cm:
        for row in table.rows:
            for cell, width in zip(row.cells, widths_cm):
                cell.width = Cm(width)

    document.add_paragraph("")


def add_figure(document: Document, image_path: Path, caption: str, width_inches: float = 6.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    cap = document.add_paragraph(caption, style="CaptionCustom")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_case_spec(document: Document, case: UseCase) -> None:
    document.add_heading(f"{case.code} {case.title}", level=4)
    meta_rows = [
        TableRow(["Version", "1.0"]),
        TableRow(["Actor principal", case.actor]),
        TableRow(["Personal involucrado e interes", case.stakeholders]),
        TableRow(["Descripcion", case.description]),
        TableRow(["Precondiciones", case.preconditions]),
        TableRow(["Postcondiciones", case.postconditions]),
        TableRow(["Extensiones o secuencia alternativa", case.alternatives]),
        TableRow(["Frecuencia esperada", case.frequency]),
    ]
    add_table(document, ["Campo", "Contenido"], meta_rows, [5.0, 11.0])
    step_rows = [TableRow([str(idx), step]) for idx, step in enumerate(case.steps, start=1)]
    add_table(document, ["Paso", "Accion"], step_rows, [2.0, 14.0])


def make_cover(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    spacer = document.add_paragraph("")
    spacer.paragraph_format.space_after = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Proyecto de Desarrollo de\nAplicaciones Multiplataforma")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb(COLOR_DARK))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("PlotSkip")
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*hex_to_rgb(COLOR_ACCENT))

    if LOGO_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(2.2))

    details = [
        "Juan Caro Vaquero",
        "Curso 2025-2026",
        "Primera version de trabajo",
        "IES Nervion",
    ]
    for item in details:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(*hex_to_rgb(COLOR_MID))

    document.add_page_break()


def build_document() -> None:
    ensure_dirs()

    context_diagram = make_context_diagram()
    use_case_diagram = make_use_case_diagram()
    architecture_diagram = make_architecture_diagram()
    analysis_class_diagram = make_analysis_class_diagram()
    design_class_diagram = make_design_class_diagram()

    document = Document()
    set_document_styles(document)
    add_header_footer(document)
    make_cover(document)

    document.add_paragraph("Indice", style="Title")
    add_toc(document.add_paragraph())
    document.add_page_break()

    document.add_heading("1. Objetivos del EVS", level=1)
    add_paragraphs(
        document,
        """
        El objetivo de PlotSkip es desarrollar una aplicacion movil centrada en el seguimiento, valoracion y conversacion social alrededor del cine y las series. La propuesta busca combinar en una sola experiencia el diario de visionado, la escritura de resenas, la curacion de listas y la relacion entre usuarios con gustos similares.

        El sistema se plantea como una red social vertical especializada en contenido audiovisual. A diferencia de soluciones generalistas o demasiado fragmentadas, PlotSkip pretende ofrecer un flujo natural para descubrir obras, registrar lo que se ha visto, opinar sobre ello y conectar con la actividad de otros usuarios desde una misma interfaz.
        """,
    )

    document.add_heading("1.1. Descripcion general del sistema", level=3)
    add_paragraphs(
        document,
        """
        PlotSkip consiste en una aplicacion movil para Android e iOS desarrollada con React Native y Expo, apoyada por una API REST construida con FastAPI y una base de datos PostgreSQL. La aplicacion permite buscar peliculas y series mediante TMDB, consultar fichas de detalle, marcar estados personales como vista o watchlist, registrar visionados por fecha y publicar resenas con puntuacion.

        El sistema dispone tambien de una capa social. Los usuarios pueden seguir a otros perfiles, ver un feed de actividad reciente, comentar resenas, votar las que consideran utiles y consultar perfiles publicos con estadisticas agregadas. A ello se suma un sistema de listas publicas o privadas, con posibilidad de colaboracion entre varios usuarios y autoria por elemento, lo que refuerza el caracter comunitario de la aplicacion.

        Desde el punto de vista de seguridad y experiencia de uso, PlotSkip utiliza autenticacion basada en JWT con refresh tokens, cierre de sesion seguro, recuperacion de contrasena por correo electronico y almacenamiento protegido del estado de sesion en el dispositivo. El resultado es una aplicacion de uso continuado, pensada para registrar habitos audiovisuales y compartirlos con otros usuarios.
        """,
    )

    document.add_heading("1.2. Diagrama de contexto", level=3)
    add_figure(document, context_diagram, "Figura 1. Diagrama de contexto de PlotSkip.", width_inches=6.6)
    add_paragraphs(
        document,
        """
        En el diagrama de contexto puede observarse que el usuario interactua exclusivamente con la aplicacion movil, mientras que la logica de negocio y la persistencia recaen en la API y la base de datos del proyecto. TMDB se emplea como proveedor externo de catalogo audiovisual, el servicio de correo se utiliza para el flujo de recuperacion de contrasena y la gestion avanzada de avatares mediante seleccion de imagen desde el movil se contempla como una evolucion natural del sistema.
        """,
    )

    document.add_heading("1.3. Estudio de la situacion actual", level=3)
    add_paragraphs(
        document,
        """
        En el mercado actual existen varias referencias parciales, pero ninguna cubre exactamente el enfoque funcional de PlotSkip. Letterboxd es el referente mas evidente para diario, resenas y listas de cine, pero no integra series con la misma naturalidad y ademas ha recibido criticas por cambios de interfaz que han ocultado acciones basicas. IMDb dispone de un catalogo enorme, pero su experiencia es menos social y mas enciclopedica. Trakt resuelve bien el seguimiento de consumo, aunque la capa comunitaria resulta menos rica para la conversacion editorial.

        A partir de ese analisis se detectan tres oportunidades. La primera es unificar peliculas y series dentro de una misma experiencia coherente. La segunda es reforzar la combinacion entre seguimiento personal y red social, de forma que el usuario no solo registre lo que ve, sino que encuentre valor en la actividad de otros perfiles. La tercera es ofrecer una arquitectura tecnica moderna y mantenible, evitando dependencias o subsistemas fuera del alcance real del producto.

        PlotSkip se posiciona, por tanto, como una solucion especializada en comunidad audiovisual, mas cercana al uso diario y a la identidad del usuario que a una simple base de datos de obras.
        """,
    )

    document.add_heading("1.4. Catalogo de requisitos", level=3)

    document.add_heading("1.4.1. Requisitos funcionales", level=4)
    functional_rows = [
        TableRow(["RF0", "Debe permitir el registro e inicio de sesion mediante email, nombre de usuario y contrasena.", "1"]),
        TableRow(["RF1", "Debe permitir la recuperacion de contrasena mediante correo electronico.", "1"]),
        TableRow(["RF2", "Debe permitir buscar peliculas y series por titulo usando TMDB.", "1"]),
        TableRow(["RF3", "Debe mostrar la ficha de detalle de una obra con metadatos y resenas.", "1"]),
        TableRow(["RF4", "Debe permitir marcar una obra como vista o como pendiente en watchlist.", "1"]),
        TableRow(["RF5", "Debe permitir registrar visionados con fecha, nota y observaciones.", "1"]),
        TableRow(["RF6", "Debe permitir crear, editar y eliminar resenas propias.", "1"]),
        TableRow(["RF7", "Debe permitir comentar y votar resenas de otros usuarios.", "2"]),
        TableRow(["RF8", "Debe permitir seguir y dejar de seguir a otros usuarios.", "1"]),
        TableRow(["RF9", "Debe mostrar un feed social con actividades recientes de los usuarios seguidos.", "1"]),
        TableRow(["RF10", "Debe permitir crear listas publicas o privadas y gestionar sus elementos.", "1"]),
        TableRow(["RF11", "Debe permitir invitar colaboradores a listas y registrar quien anadio cada obra.", "2"]),
        TableRow(["RF12", "Debe mostrar perfil propio y publico con estadisticas de consumo.", "1"]),
        TableRow(["RF13", "Debe permitir editar los datos del perfil y mantener una sesion persistente con refresh token.", "1"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Prioridad"], functional_rows, [2.1, 12.9, 2.0])

    document.add_heading("1.4.2. Requisitos de datos", level=4)
    data_rows = [
        TableRow(["RD0", "Deben almacenarse usuarios, credenciales cifradas y metadatos de perfil.", "1"]),
        TableRow(["RD1", "Deben almacenarse refresh tokens y tokens de recuperacion con expiracion.", "1"]),
        TableRow(["RD2", "Deben almacenarse resenas, comentarios y votos sobre resenas.", "1"]),
        TableRow(["RD3", "Deben almacenarse estados personales de media y entradas del diario de visionado.", "1"]),
        TableRow(["RD4", "Deben almacenarse relaciones de seguimiento entre usuarios.", "1"]),
        TableRow(["RD5", "Deben almacenarse listas, colaboradores, invitaciones y autoria de cada elemento.", "1"]),
        TableRow(["RD6", "Debe mantenerse una cache minima de metadatos de TMDB para acelerar vistas y feeds.", "2"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Prioridad"], data_rows, [2.1, 12.9, 2.0])

    document.add_heading("1.4.3. Requisitos de interfaz", level=4)
    interface_rows = [
        TableRow(["RI0", "Debe existir una navegacion inferior clara hacia Inicio, Social, Diario, Listas y Perfil.", "1"]),
        TableRow(["RI1", "Debe mostrarse la informacion principal de una obra sin ocultar acciones basicas.", "1"]),
        TableRow(["RI2", "Deben mostrarse estados de carga, error y vacio de forma comprensible.", "1"]),
        TableRow(["RI3", "Debe mostrarse el perfil publico con datos sociales y estadisticos.", "2"]),
        TableRow(["RI4", "Debe permitirse la edicion simple del perfil propio sin romper la experiencia de navegacion.", "2"]),
        TableRow(["RI5", "Debe distinguirse visualmente si una lista es propia, compartida o solo de lectura.", "2"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Prioridad"], interface_rows, [2.1, 12.9, 2.0])

    document.add_heading("1.4.4. Requisitos no funcionales", level=4)
    non_functional_rows = [
        TableRow(["RNF0", "Debe funcionar de forma correcta en dispositivos Android e iOS.", "1"]),
        TableRow(["RNF1", "Debe ofrecer tiempos de respuesta razonables en operaciones frecuentes.", "2"]),
        TableRow(["RNF2", "Debe almacenar contrasenas mediante hash y no en texto plano.", "1"]),
        TableRow(["RNF3", "Debe usar HTTPS y autenticacion basada en tokens firmados.", "1"]),
        TableRow(["RNF4", "Debe ser mantenible mediante separacion de capas y contratos explicitos.", "1"]),
        TableRow(["RNF5", "Debe tolerar fallos temporales de TMDB degradando la experiencia sin bloquear la app.", "2"]),
        TableRow(["RNF6", "Debe facilitar testing automatizado de backend y validacion de tipos en mobile.", "2"]),
        TableRow(["RNF7", "Debe evitar introducir dependencias o subsistemas ajenos al alcance funcional real del proyecto.", "3"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Prioridad"], non_functional_rows, [2.1, 12.9, 2.0])

    document.add_heading("1.5. Alternativas a la solucion", level=3)
    document.add_heading("1.5.1. Alternativa I: React Native + FastAPI + PostgreSQL", level=4)
    add_paragraphs(
        document,
        """
        La primera alternativa propone un frontend movil en React Native con Expo y un backend en Python mediante FastAPI. La base de datos se apoya en PostgreSQL y el catalogo audiovisual se integra a traves de TMDB. La autenticacion se resuelve con JWT y refresh tokens, mientras que la persistencia local del cliente se protege con SecureStore.
        """,
    )
    document.add_heading("1.5.1.1. Requisitos", level=5)
    add_bullets(
        document,
        [
            "Despliegue orientado a Android e iOS.",
            "Backend desplegable en un servicio cloud con soporte Python y PostgreSQL.",
            "Integracion de TMDB y correo electronico para recuperacion de contrasena.",
        ],
    )
    document.add_heading("1.5.1.2. Tecnologias implementadas", level=5)
    add_bullets(
        document,
        [
            "React Native con Expo para la aplicacion movil.",
            "FastAPI y Pydantic para la API REST.",
            "PostgreSQL con SQLAlchemy 2.0 y Alembic.",
            "JWT con refresh tokens y almacenamiento seguro de sesion.",
            "TMDB como proveedor de catalogo de peliculas y series.",
            "Resend para el envio de correos de recuperacion.",
        ],
    )
    document.add_heading("1.5.1.3. Estudio de riesgo", level=5)
    alt1_risk_rows = [
        TableRow(["Dependencia de TMDB", "20%", "1"]),
        TableRow(["Complejidad de integracion mobile-backend", "25%", "1"]),
        TableRow(["Desajustes de UX entre plataformas", "15%", "2"]),
        TableRow(["Sobrecoste de mantenimiento de API propia", "15%", "2"]),
        TableRow(["Curva de aprendizaje de Expo Router y arquitectura limpia", "20%", "2"]),
    ]
    add_table(document, ["Riesgo", "Porcentaje", "Prioridad"], alt1_risk_rows, [9.5, 3.0, 3.0])

    document.add_heading("1.5.2. Alternativa II: Flutter + Node.js + MongoDB", level=4)
    add_paragraphs(
        document,
        """
        La segunda alternativa plantea una app movil en Flutter con backend JavaScript sobre Node.js y almacenamiento principal en MongoDB. La propuesta permitiria una base tecnologica unificada alrededor de JavaScript y Dart, con una capa de datos mas flexible, aunque con menor afinidad hacia el modelo relacional del dominio.
        """,
    )
    document.add_heading("1.5.2.1. Requisitos", level=5)
    add_bullets(
        document,
        [
            "Despliegue movil para Android e iOS.",
            "Backend en Node.js con soporte para autenticacion y consumo de APIs externas.",
            "Modelo de datos con validaciones adicionales para evitar inconsistencias.",
        ],
    )
    document.add_heading("1.5.2.2. Tecnologias implementadas", level=5)
    add_bullets(
        document,
        [
            "Flutter para la aplicacion cliente.",
            "Node.js con Express para la API.",
            "MongoDB como almacenamiento principal.",
            "JWT para autenticacion.",
            "TMDB como fuente de metadatos audiovisuales.",
        ],
    )
    document.add_heading("1.5.2.3. Estudio de riesgo", level=5)
    alt2_risk_rows = [
        TableRow(["Inconsistencia de datos en modelo flexible", "30%", "1"]),
        TableRow(["Mayor esfuerzo en joins logicos del dominio social", "25%", "1"]),
        TableRow(["Complejidad adicional en agregaciones de estadisticas", "20%", "2"]),
        TableRow(["Mantenimiento de dos ecosistemas distintos", "15%", "2"]),
        TableRow(["Menor alineacion con el backend ya construido", "35%", "1"]),
    ]
    add_table(document, ["Riesgo", "Porcentaje", "Prioridad"], alt2_risk_rows, [9.5, 3.0, 3.0])

    document.add_heading("1.6. Alternativa seleccionada y justificacion", level=3)
    add_paragraphs(
        document,
        """
        Se selecciona la primera alternativa, formada por React Native con Expo, FastAPI y PostgreSQL, por ser la que mejor se alinea con el estado actual del proyecto y con la naturaleza relacional del dominio. PlotSkip no necesita una base de datos documental ni un backend improvisado alrededor de tiempo real generalista; necesita consistencia en usuarios, follow, resenas, listas y estadisticas.

        FastAPI permite mantener una API clara, tipada y facil de probar, mientras que PostgreSQL aporta integridad referencial y consultas expresivas para el feed social, el diario, las listas colaborativas y las estadisticas agregadas. En el cliente, React Native con Expo ofrece una iteracion rapida sobre iOS y Android sin comprometer la experiencia movil. La eleccion final no responde solo a preferencias tecnologicas, sino a una mejor correspondencia entre problema, arquitectura y madurez real del desarrollo.
        """,
    )

    document.add_heading("2. Gestion del proyecto", level=1)
    add_paragraphs(
        document,
        """
        Para el desarrollo del proyecto se ha seguido una planificacion incremental inspirada en Scrum, utilizando un backlog vivo y sprints funcionales centrados en entregables de valor. La prioridad no ha sido construir todas las ideas imaginables, sino consolidar una vertical de producto coherente: autenticacion, catalogo, diario, resenas, social, listas y perfil.
        """,
    )

    document.add_heading("2.1. Backlog", level=3)
    backlog_rows = [
        TableRow(["Auth", "LG0", "Como usuario nuevo, quiero registrarme con email, username y contrasena para crear mi cuenta."]),
        TableRow(["Auth", "LG1", "Como usuario registrado, quiero iniciar sesion y mantenerla entre usos de la app."]),
        TableRow(["Auth", "LG2", "Como usuario, quiero recuperar mi contrasena por correo si la olvido."]),
        TableRow(["Auth", "LG3", "Como usuario, quiero cerrar sesion invalidando mi refresh token."]),
        TableRow(["Descubrimiento", "DS0", "Como usuario, quiero buscar peliculas y series por titulo."]),
        TableRow(["Descubrimiento", "DS1", "Como usuario, quiero explorar contenido destacado desde la pantalla de inicio."]),
        TableRow(["Detalle", "DT0", "Como usuario, quiero ver una ficha editorial de cada obra con metadata y resenas."]),
        TableRow(["Detalle", "DT1", "Como usuario, quiero marcar una obra como vista o watchlist desde su detalle."]),
        TableRow(["Diario", "DV0", "Como usuario, quiero registrar el dia exacto de visionado y mi puntuacion."]),
        TableRow(["Diario", "DV1", "Como usuario, quiero repetir visionados de una misma obra en fechas distintas."]),
        TableRow(["Resenas", "RS0", "Como usuario, quiero escribir una resena con nota y marca de spoilers."]),
        TableRow(["Resenas", "RS1", "Como usuario, quiero editar o eliminar mis resenas."]),
        TableRow(["Resenas", "RS2", "Como usuario, quiero comentar y votar resenas de otros usuarios."]),
        TableRow(["Social", "SC0", "Como usuario, quiero seguir a otros perfiles para ver su actividad."]),
        TableRow(["Social", "SC1", "Como usuario, quiero consultar un feed con follow, listas, resenas y visionados."]),
        TableRow(["Social", "SC2", "Como usuario, quiero buscar perfiles por nombre de usuario."]),
        TableRow(["Listas", "LS0", "Como usuario, quiero crear listas publicas o privadas de obras."]),
        TableRow(["Listas", "LS1", "Como usuario, quiero reordenar y gestionar los elementos de mis listas."]),
        TableRow(["Listas", "LS2", "Como usuario, quiero compartir listas con colaboradores de confianza."]),
        TableRow(["Perfil", "PF0", "Como usuario, quiero editar mi perfil y mi biografia."]),
        TableRow(["Perfil", "PF1", "Como usuario, quiero ver mis estadisticas de consumo y favoritas."]),
        TableRow(["Perfil", "PF2", "Como usuario, quiero consultar el perfil publico de otros usuarios."]),
        TableRow(["Evolucion", "EV0", "Como usuario, quiero poder seleccionar mi foto de perfil directamente desde el movil."]),
    ]
    add_table(document, ["Tema", "ID", "User story"], backlog_rows, [3.0, 2.0, 11.0])

    document.add_heading("2.2. Sprints", level=3)
    sprint_rows = [
        TableRow(["Primer sprint", "Base tecnica, registro, login y persistencia de sesion.", "18/03/2025 - 29/03/2025", "11 dias"]),
        TableRow(["Segundo sprint", "Busqueda de media, detalle y estados personales de visionado.", "30/03/2025 - 12/04/2025", "14 dias"]),
        TableRow(["Tercer sprint", "Resenas, comentarios y primeras capacidades sociales.", "13/04/2025 - 26/04/2025", "14 dias"]),
        TableRow(["Cuarto sprint", "Diario de visionado, feed social y perfiles publicos.", "27/04/2025 - 11/05/2025", "15 dias"]),
        TableRow(["Quinto sprint", "Listas, estadisticas, favoritas y perfil curado.", "12/05/2025 - 31/05/2025", "20 dias"]),
        TableRow(["Sexto sprint", "Listas colaborativas, pulido de UX y preparacion de la memoria.", "01/06/2025 - 12/06/2025", "12 dias"]),
    ]
    add_table(document, ["Sprint", "Objetivo", "Fechas", "Duracion"], sprint_rows, [4.0, 8.4, 3.2, 2.2])
    add_paragraphs(
        document,
        """
        Esta distribucion recoge la evolucion funcional del proyecto de una forma coherente con los documentos internos y con el estado del repositorio. La planificacion ha sido adaptativa: algunas decisiones de producto se han redefinido conforme avanzaba el desarrollo para mantener el foco en la experiencia audiovisual y social realmente perseguida.
        """,
    )

    document.add_heading("3. Analisis de Sistemas de informacion (ASI)", level=1)
    document.add_heading("3.1. Descripcion general del entorno tecnologico del sistema", level=3)
    add_paragraphs(
        document,
        """
        En la fase de analisis se identifica PlotSkip como un sistema cliente-servidor compuesto por una aplicacion movil y una API REST. La app movil se desarrolla en TypeScript sobre React Native y Expo, y la API se construye con FastAPI siguiendo una separacion clara entre dominio, datos, presentacion e infraestructura.

        El sistema utiliza PostgreSQL como fuente principal de verdad para usuarios, relaciones sociales, resenas, diario y listas. La informacion audiovisual se consulta en TMDB y se normaliza en una cache minima para acelerar listados, feeds y enriquecimiento de estadisticas. La sesion de usuario se gestiona con access token de corta vida y refresh token persistido, lo que reduce la friccion al reabrir la aplicacion.

        El alcance productivo se concentra en Android e iOS. Aunque Expo permita ejecutar una version web de desarrollo, esa capacidad no forma parte del objetivo funcional de este proyecto ni de su documentacion final como producto.
        """,
    )

    document.add_heading("3.2. Catalogo de usuarios", level=3)
    user_rows = [
        TableRow(["USUV0", "Visitante no autenticado", "Puede registrarse, iniciar sesion y recuperar contrasena."]),
        TableRow(["USUV1", "Usuario registrado", "Puede buscar contenido, registrar visionados, escribir resenas, seguir usuarios y crear listas."]),
        TableRow(["USUV2", "Propietario de lista", "Puede editar una lista propia, gestionar su privacidad e invitar colaboradores."]),
        TableRow(["USUV3", "Colaborador de lista", "Puede editar una lista compartida aceptada y anadir o eliminar elementos segun permisos."]),
    ]
    add_table(document, ["Codigo", "Tipo de usuario", "Responsabilidad principal"], user_rows, [2.3, 4.2, 9.5])

    document.add_heading("3.3. Modelo de casos de uso", level=3)
    add_figure(document, use_case_diagram, "Figura 2. Modelo resumido de casos de uso principales.", width_inches=6.8)
    add_paragraphs(
        document,
        """
        El modelo de casos de uso se centra en un unico actor principal, el usuario de PlotSkip, que atraviesa distintos momentos de la experiencia: acceso al sistema, descubrimiento de contenido, registro de actividad, conversacion social y gestion de identidad. Algunas acciones generan estados derivados, como el feed social o las estadisticas, pero el origen sigue siendo la actividad voluntaria del usuario dentro de la aplicacion.
        """,
    )

    document.add_heading("3.4. Especificacion de casos de uso", level=3)
    use_cases = [
        UseCase(
            "CU1",
            "Registrarse",
            "USUV0",
            "USUV0, sistema",
            "El sistema debe permitir la creacion de una cuenta nueva con identificadores unicos y credenciales validas.",
            "El usuario no debe existir previamente con el mismo email o username.",
            "Se crea el usuario y puede iniciar sesion posteriormente.",
            [
                "El usuario abre la pantalla de registro.",
                "Introduce username, email, contrasena y confirmacion de contrasena.",
                "Pulsa el boton de registro.",
                "El sistema valida el formato y la unicidad de los datos.",
                "El sistema crea la cuenta y devuelve la respuesta correspondiente.",
            ],
            "Si el email o el username ya existen, el sistema rechaza la operacion.",
            "20/dia",
        ),
        UseCase(
            "CU2",
            "Iniciar sesion",
            "USUV0",
            "USUV0, sistema",
            "El sistema autentica al usuario y entrega tokens de acceso y refresco.",
            "La cuenta del usuario debe existir y la contrasena debe ser correcta.",
            "El usuario queda autenticado en la aplicacion.",
            [
                "El usuario abre la pantalla de login.",
                "Introduce email y contrasena.",
                "Pulsa iniciar sesion.",
                "El sistema valida las credenciales.",
                "El sistema entrega los tokens y la app almacena la sesion.",
            ],
            "Si la contrasena es incorrecta, el sistema devuelve error de autenticacion.",
            "60/dia",
        ),
        UseCase(
            "CU3",
            "Recuperar contrasena",
            "USUV0",
            "USUV0, sistema, servicio de correo",
            "El sistema permite solicitar un enlace o flujo de restablecimiento de contrasena.",
            "El email debe pertenecer a una cuenta existente.",
            "Se emite un token temporal y el usuario puede establecer una nueva contrasena.",
            [
                "El usuario accede a la opcion de recuperar contrasena.",
                "Introduce su correo electronico.",
                "El sistema genera un token seguro con caducidad.",
                "Se envia el correo con las instrucciones de reseteo.",
                "El usuario establece una nueva contrasena desde el flujo de restablecimiento.",
            ],
            "Si el correo no existe, el sistema responde de forma controlada para no exponer informacion sensible.",
            "5/dia",
        ),
        UseCase(
            "CU4",
            "Buscar contenido",
            "USUV1",
            "USUV1, sistema, TMDB",
            "El usuario localiza peliculas o series a partir de un texto de busqueda.",
            "El usuario debe estar autenticado.",
            "Se muestra una lista de resultados relevantes.",
            [
                "El usuario escribe un termino de busqueda.",
                "El sistema consulta TMDB y/o la cache local.",
                "Se presentan resultados con poster, titulo y ano.",
            ],
            "Si TMDB no responde, el sistema muestra un error y permite reintentar.",
            "120/dia",
        ),
        UseCase(
            "CU5",
            "Registrar visionado",
            "USUV1",
            "USUV1, sistema",
            "El usuario anade una entrada al diario con fecha, nota y observaciones.",
            "La obra debe existir en el contexto de la app.",
            "Se crea una nueva entrada en watch_log y la obra queda marcada como vista.",
            [
                "El usuario abre el detalle de una obra.",
                "Indica fecha de visionado, nota y notas opcionales.",
                "Pulsa guardar.",
                "El sistema crea el registro y actualiza los estados personales.",
            ],
            "Si la fecha es futura o la nota esta fuera de rango, el sistema rechaza la operacion.",
            "80/dia",
        ),
        UseCase(
            "CU6",
            "Publicar resena",
            "USUV1",
            "USUV1, sistema",
            "El usuario escribe una resena con puntuacion y marca opcional de spoilers.",
            "Debe existir una obra seleccionada y el usuario debe estar autenticado.",
            "La resena queda publicada y puede aparecer en feed y detalle de media.",
            [
                "El usuario abre el modal o formulario de resena desde la ficha de la obra.",
                "Escribe el titulo opcional, el cuerpo y la nota.",
                "Marca si contiene spoilers.",
                "Pulsa publicar.",
                "El sistema persiste la resena y genera la actividad asociada.",
            ],
            "Si el usuario ya tenia una resena sobre esa obra, el sistema debe editarla o rechazar duplicados segun el flujo.",
            "45/dia",
        ),
        UseCase(
            "CU7",
            "Seguir usuario",
            "USUV1",
            "USUV1, sistema",
            "El usuario decide seguir a otro perfil para ver su actividad.",
            "No se permite seguirse a si mismo.",
            "La relacion de follow queda creada y alimenta el feed social.",
            [
                "El usuario abre un perfil publico o un resultado de busqueda.",
                "Pulsa el boton de seguir.",
                "El sistema registra la relacion y actualiza contadores.",
            ],
            "Si la relacion ya existia, la operacion es idempotente.",
            "70/dia",
        ),
        UseCase(
            "CU8",
            "Consultar feed social",
            "USUV1",
            "USUV1, sistema",
            "El sistema muestra actividades recientes de los usuarios seguidos.",
            "El usuario debe seguir al menos a otro usuario para obtener contenido relevante.",
            "Se muestra un feed paginado por cursor con actividades cronologicas.",
            [
                "El usuario abre la tab de Inicio o Social.",
                "El sistema consulta el feed de actividades de usuarios seguidos.",
                "Se presentan resenas, follow, listas y visionados recientes.",
            ],
            "Si el usuario no sigue a nadie, se muestra un estado vacio orientado al descubrimiento.",
            "100/dia",
        ),
        UseCase(
            "CU9",
            "Crear lista",
            "USUV1",
            "USUV1, sistema",
            "El usuario crea una lista tematica de peliculas o series.",
            "El usuario debe estar autenticado.",
            "La lista queda almacenada como publica o privada segun la configuracion elegida.",
            [
                "El usuario accede a la seccion de listas.",
                "Pulsa crear lista.",
                "Introduce nombre, descripcion y visibilidad.",
                "El sistema crea la lista y la devuelve al cliente.",
            ],
            "Si el nombre esta vacio, la lista no debe crearse.",
            "20/dia",
        ),
        UseCase(
            "CU10",
            "Invitar colaborador a lista",
            "USUV2",
            "USUV2, sistema, usuario invitado",
            "El propietario de una lista comparte su mantenimiento con otro usuario.",
            "Debe existir follow mutuo entre ambas cuentas y la lista debe pertenecer al emisor.",
            "Se genera una invitacion pendiente que el otro usuario podra aceptar o rechazar.",
            [
                "El propietario abre el detalle de una lista propia.",
                "Busca usuarios elegibles por follow mutuo.",
                "Selecciona al usuario e invita.",
                "El sistema crea la invitacion y la deja pendiente.",
            ],
            "No se puede invitar dos veces al mismo usuario ni invitar a un colaborador ya activo.",
            "8/dia",
        ),
        UseCase(
            "CU11",
            "Aceptar invitacion a lista",
            "USUV3",
            "USUV3, sistema",
            "El usuario acepta una colaboracion y pasa a editar una lista compartida.",
            "Debe existir una invitacion pendiente dirigida al usuario.",
            "La lista pasa a figurar entre las compartidas y el usuario adquiere permisos de edicion.",
            [
                "El usuario entra en su area de listas compartidas o invitaciones.",
                "Pulsa aceptar invitacion.",
                "El sistema activa la colaboracion y actualiza la relacion del usuario con la lista.",
            ],
            "Si la invitacion ya fue resuelta, la accion no debe aplicarse otra vez.",
            "5/dia",
        ),
        UseCase(
            "CU12",
            "Consultar perfil y estadisticas",
            "USUV1",
            "USUV1, sistema",
            "El usuario consulta su propio perfil o el de otro usuario para conocer actividad y gustos.",
            "El usuario debe estar autenticado.",
            "Se muestran datos de perfil, contadores sociales y estadisticas agregadas desde watch_log.",
            [
                "El usuario abre su perfil o el de otro usuario.",
                "El sistema consulta informacion publica o privada segun corresponda.",
                "Se muestran biografia, contadores, favoritas y estadisticas de visionado.",
            ],
            "Si fallan las llamadas de enriquecimiento TMDB, el sistema debe degradar sin bloquear la pantalla completa.",
            "90/dia",
        ),
    ]
    for case in use_cases:
        add_case_spec(document, case)

    document.add_heading("3.5. Modelo de clases de analisis", level=3)
    add_figure(document, analysis_class_diagram, "Figura 3. Modelo de clases de analisis.", width_inches=6.9)
    entity_rows = [
        TableRow(["Usuario", "Representa a la persona autenticada en la plataforma y concentra identidad, perfil y relaciones sociales."]),
        TableRow(["MediaRef", "Referencia ligera a una obra de TMDB mediante tmdb_id y media_type."]),
        TableRow(["Resena", "Opinion editorial de un usuario sobre una obra concreta."]),
        TableRow(["WatchLog", "Registro temporal de visionado, con soporte para rewatch."]),
        TableRow(["Lista", "Coleccion tematica de obras con privacidad configurable."]),
        TableRow(["ElementoLista", "Relacion ordenada entre lista y obra, con autoria por usuario."]),
        TableRow(["Comentario", "Interaccion textual sobre una resena."]),
        TableRow(["Follow", "Relacion social dirigida entre dos usuarios."]),
    ]
    add_table(document, ["Entidad", "Descripcion"], entity_rows, [4.0, 12.0])

    document.add_heading("3.6. Interfaces de usuario", level=3)
    document.add_heading("3.6.1. Aspectos comunes de la interfaz de usuario", level=4)
    add_bullets(
        document,
        [
            "Barra de navegacion inferior con cinco accesos principales: Inicio, Social, Diario, Listas y Perfil.",
            "Jerarquia visual clara para destacar posters, titulos, valoraciones y acciones frecuentes.",
            "Estados de carga y vacio visibles para evitar pantallas ambiguas.",
            "Composicion editorial en pantallas de detalle, inspirada en patrones de apps audiovisuales modernas.",
            "Uso de componentes reutilizables y estilos consistentes entre formularios, cards y botones.",
        ],
    )

    document.add_heading("3.6.2. Especificacion de pantallas y ventanas", level=4)
    screen_rows = [
        TableRow(["Login", "Pantalla de acceso con email, contrasena y enlace a recuperacion."]),
        TableRow(["Registro", "Pantalla de alta con confirmacion de contrasena para reducir errores."]),
        TableRow(["Inicio", "Pantalla de descubrimiento y acceso rapido a actividad social resumida."]),
        TableRow(["Social", "Feed visual con actividad de usuarios seguidos y acceso a perfiles."]),
        TableRow(["Detalle de obra", "Ficha editorial con poster, metadata, estado personal y resenas."]),
        TableRow(["Diario", "Listado cronologico de visionados personales."]),
        TableRow(["Listas", "Area de listas propias, compartidas e invitaciones pendientes."]),
        TableRow(["Detalle de lista", "Vista completa de una lista con permisos, colaboradores y autoria por elemento."]),
        TableRow(["Perfil propio", "Pantalla de identidad, biografia, stats, favoritas y visionados recientes."]),
        TableRow(["Perfil publico", "Perfil de otro usuario con follow, stats y resumen de actividad."]),
        TableRow(["Busqueda de usuarios", "Pantalla auxiliar para localizar y abrir perfiles."]),
    ]
    add_table(document, ["Pantalla", "Descripcion funcional"], screen_rows, [4.5, 11.5])

    document.add_heading("4. Diseno de Sistemas de informacion (DSI)", level=1)
    document.add_heading("4.1. Diseno de la arquitectura del sistema", level=3)
    document.add_heading("4.1.1. Descripcion general del entorno tecnologico del sistema", level=4)
    add_figure(document, architecture_diagram, "Figura 4. Arquitectura general del sistema.", width_inches=6.8)
    add_paragraphs(
        document,
        """
        La arquitectura del sistema separa con claridad el cliente movil, el backend y los servicios externos. La app movil concentra la experiencia de usuario y delega en la API la autenticacion, la persistencia del dominio y la coordinacion de reglas de negocio. La API integra TMDB para metadatos audiovisuales, PostgreSQL para almacenamiento relacional y un servicio de correo para la recuperacion de contrasena.

        En el cliente se utiliza almacenamiento seguro para tokens y una capa HTTP con logica de refresco transparente de sesion. En el servidor, la estructura en capas evita acoplar la presentacion a la persistencia, lo que facilita la prueba automatizada y el crecimiento del proyecto.
        """,
    )

    document.add_heading("4.1.2. Catalogo de requisitos de diseno", level=4)
    design_req_rows = [
        TableRow(["RD0", "La arquitectura debe separar dominio, datos, presentacion e infraestructura.", "18/03/2025"]),
        TableRow(["RD1", "La base de datos principal debe responder a un modelo relacional.", "18/03/2025"]),
        TableRow(["RD2", "La autenticacion debe soportar access token y refresh token.", "20/03/2025"]),
        TableRow(["RD3", "La API debe permitir testing automatizado de sus contratos principales.", "20/03/2025"]),
        TableRow(["RD4", "La app debe degradar de forma segura cuando fallen servicios externos como TMDB.", "22/03/2025"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Fecha"], design_req_rows, [2.1, 11.9, 2.0])
    view_req_rows = [
        TableRow(["RV0", "La interfaz debe ser comprensible y estable para el usuario habitual.", "24/03/2025"]),
        TableRow(["RV1", "Las acciones principales no deben quedar ocultas tras varios pasos.", "24/03/2025"]),
        TableRow(["RV2", "Los formularios deben validar entrada antes de llegar al backend cuando sea posible.", "24/03/2025"]),
        TableRow(["RV3", "La experiencia debe personalizarse con estados, favoritas y actividad del usuario.", "26/03/2025"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Fecha"], view_req_rows, [2.1, 11.9, 2.0])

    document.add_heading("4.2. Modelo de las clases de diseno", level=3)
    add_figure(document, design_class_diagram, "Figura 5. Modelo de diseno por capas.", width_inches=6.8)
    design_class_rows = [
        TableRow(["Presentacion", "LoginScreen, DetailScreen, ProfileScreen, ListDetailScreen, SocialScreen y sus ViewModels."]),
        TableRow(["Dominio", "Entidades User, Review, WatchLog, List y contratos de repositorio."]),
        TableRow(["Datos", "Repositorios HTTP en mobile y SQLAlchemy en backend; mapeadores y servicios externos."]),
        TableRow(["Infraestructura", "Configuracion, sesion de base de datos, JWT helpers, token storage y email."]),
    ]
    add_table(document, ["Capa", "Elementos de diseno representativos"], design_class_rows, [4.2, 11.8])

    document.add_heading("4.3. Modelo fisico de datos", level=3)
    add_paragraphs(
        document,
        """
        El modelo fisico de datos de PlotSkip se implementa sobre PostgreSQL y se apoya en tablas con claves foraneas, restricciones de unicidad e indices para las consultas mas frecuentes. La decision prioriza consistencia, trazabilidad y facilidad de consulta en un dominio claramente relacional.
        """,
    )
    physical_rows = [
        TableRow(["users", "id, email, username, password_hash, display_name, bio, avatar_url, created_at, updated_at", "Entidad base de identidad."]),
        TableRow(["refresh_tokens", "user_id, token_hash, expires_at, created_at", "Persistencia de sesion renovable."]),
        TableRow(["password_reset_tokens", "user_id, token_hash, expires_at, used", "Recuperacion de contrasena."]),
        TableRow(["follows", "follower_id, followed_id, created_at", "Relacion social entre usuarios."]),
        TableRow(["user_media_status", "user_id, tmdb_id, media_type, status, created_at", "Estados personales de vista/watchlist."]),
        TableRow(["watch_log", "user_id, tmdb_id, media_type, watched_at, rating, notes, created_at", "Diario con soporte de rewatch."]),
        TableRow(["reviews", "user_id, tmdb_id, media_type, rating, title, body, contains_spoilers, created_at, updated_at", "Resenas por obra y usuario."]),
        TableRow(["comments", "review_id, user_id, body, created_at", "Comentarios sobre resenas."]),
        TableRow(["review_votes", "review_id, user_id, created_at", "Votos de utilidad."]),
        TableRow(["lists", "user_id, name, description, is_public, created_at, updated_at", "Cabecera de listas."]),
        TableRow(["list_items", "list_id, tmdb_id, media_type, position, added_at, added_by_user_id", "Elementos ordenados y con autoria."]),
        TableRow(["list_collaborators", "list_id, user_id, created_at", "Usuarios activos con permisos de colaboracion."]),
        TableRow(["list_invitations", "list_id, invitee_user_id, invited_by_user_id, status, created_at, responded_at", "Invitaciones pendientes o resueltas."]),
        TableRow(["activities", "user_id, activity_type, reference_id, tmdb_id, media_type, created_at", "Base del feed social."]),
        TableRow(["user_favorite_media", "user_id, position, tmdb_id, media_type", "Favoritas manuales de perfil."]),
    ]
    add_table(document, ["Tabla", "Campos principales", "Finalidad"], physical_rows, [3.8, 8.7, 4.0])

    document.add_heading("4.4. Diseno de la interfaz de usuario", level=3)
    add_paragraphs(
        document,
        """
        El diseno de interfaz prioriza tres ideas: accion visible, lectura rapida y continuidad entre pantallas. La aplicacion organiza la experiencia principal en cinco tabs. Inicio concentra descubrimiento y acceso al contenido destacado; Social presenta el feed de actividad; Diario recoge los visionados; Listas organiza la curacion y la colaboracion; y Perfil resume identidad y estadisticas.

        La ficha de detalle adopta una composicion editorial, con protagonismo del poster, metadatos resumidos y tres acciones principales faciles de localizar: marcar vista, guardar en watchlist y escribir resena. En listas se diferencia claramente entre listas propias, compartidas e invitaciones pendientes. En perfil se muestran biografia, favoritas y estadisticas, evitando que la pantalla se convierta en un simple menu de configuracion.

        Como evolucion prevista, el sistema de avatar pasara de una URL segura editable a una seleccion directa de imagen desde archivos del dispositivo movil, con posterior subida a almacenamiento externo.
        """,
    )

    document.add_heading("4.5. Plan de migracion y carga inicial de datos", level=3)
    add_paragraphs(
        document,
        """
        PlotSkip no migra informacion desde una aplicacion anterior, por lo que no existe una carga historica compleja de negocio. La puesta en marcha del sistema se basa en tres pasos: configuracion de entorno, aplicacion de migraciones de base de datos y conexion con servicios externos como TMDB y correo electronico.

        En la primera ejecucion no se precargan usuarios ni catalogos completos. El catalogo audiovisual se resuelve bajo demanda a partir de TMDB, con una cache local de datos basicos para acelerar consultas repetidas. De esta forma se reduce el volumen de datos inicial, se simplifica el mantenimiento y se evita duplicar innecesariamente informacion externa.
        """,
    )

    document.add_heading("4.6. Plan de pruebas", level=3)
    add_paragraphs(
        document,
        """
        El proyecto combina pruebas automatizadas de backend con validaciones de tipado en mobile y revision funcional de las pantallas clave. El objetivo del plan de pruebas no es solo comprobar que cada endpoint responde, sino garantizar que se cumplen reglas de negocio importantes como la idempotencia de follow, la privacidad de listas o el soporte de rewatch en el diario.
        """,
    )

    document.add_heading("4.6.1. Pruebas del login, registro y presentacion", level=4)
    auth_test_rows = [
        TableRow(["PR-LRP0", "Inicio de sesion correcto de un usuario registrado.", "V"]),
        TableRow(["PR-LRP1", "Inicio de sesion con contrasena incorrecta.", "F"]),
        TableRow(["PR-LRP2", "Inicio de sesion con email desconocido.", "F"]),
        TableRow(["PR-LRP3", "Recuperacion de contrasena para usuario existente.", "V"]),
        TableRow(["PR-LRP4", "Obtencion de perfil mediante /auth/me con campos ampliados.", "V"]),
        TableRow(["PR-LRP5", "Persistencia y renovacion transparente de sesion mediante refresh token.", "V"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Tipo"], auth_test_rows, [3.0, 11.0, 2.0])

    document.add_heading("4.6.2. Pruebas de descubrimiento, detalle y diario", level=4)
    discovery_test_rows = [
        TableRow(["PR-D0", "Busqueda de contenido por titulo.", "V"]),
        TableRow(["PR-D1", "Marcado de estado watchlist.", "V"]),
        TableRow(["PR-D2", "Creacion de watch log con nota valida.", "V"]),
        TableRow(["PR-D3", "Creacion de watch log sin nota.", "V"]),
        TableRow(["PR-D4", "Rechazo de fecha futura en watch log.", "F"]),
        TableRow(["PR-D5", "Listado de rewatch ordenado por fecha descendente.", "V"]),
        TableRow(["PR-D6", "Enriquecimiento de visionados recientes con resumen de media.", "V"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Tipo"], discovery_test_rows, [3.0, 11.0, 2.0])

    document.add_heading("4.6.3. Pruebas de social, resenas y listas", level=4)
    social_test_rows = [
        TableRow(["PR-S0", "Busqueda de usuarios mostrando estado de follow.", "V"]),
        TableRow(["PR-S1", "Follow idempotente y rechazo de self-follow.", "V/F"]),
        TableRow(["PR-S2", "Feed con actividad solo de usuarios seguidos y paginacion por cursor.", "V"]),
        TableRow(["PR-S3", "Actualizacion y lectura de favoritas de perfil.", "V"]),
        TableRow(["PR-S4", "Creacion de lista y lectura desde /lists/me.", "V"]),
        TableRow(["PR-S5", "Lista privada oculta a terceros y visible a colaborador aceptado.", "V"]),
        TableRow(["PR-S6", "Invitacion a colaboracion requiriendo follow mutuo.", "V/F"]),
        TableRow(["PR-S7", "Colaborador aceptado puede editar metadatos y elementos.", "V"]),
        TableRow(["PR-S8", "Owner puede retirar a un colaborador.", "V"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Tipo"], social_test_rows, [3.0, 11.0, 2.0])

    document.add_heading("4.6.4. Pruebas de configuracion y robustez", level=4)
    config_test_rows = [
        TableRow(["PR-CF0", "Actualizacion de display_name y bio normalizando espacios.", "V"]),
        TableRow(["PR-CF1", "Rechazo de avatar_url insegura.", "F"]),
        TableRow(["PR-CF2", "Rechazo de display_name vacio.", "F"]),
        TableRow(["PR-CF3", "Stats vacias para usuario sin watch log.", "V"]),
        TableRow(["PR-CF4", "Stats con agregacion real de watch log y degradacion si falla TMDB.", "V"]),
        TableRow(["PR-CF5", "Visual feed agrupado por media reciente.", "V"]),
    ]
    add_table(document, ["Codigo", "Descripcion", "Tipo"], config_test_rows, [3.0, 11.0, 2.0])

    note_heading = document.add_heading("Observaciones de esta primera version", level=2)
    note_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_bullets(
        document,
        [
            "Se ha depurado el contenido heredado para dejar una memoria alineada con el alcance y el estado real de PlotSkip.",
            "La estructura reproduce los apartados del documento de referencia, pero ajustados al dominio real de PlotSkip.",
            "Los diagramas incluidos son funcionales y suficientes para una primera version, aunque pueden refinarse visualmente en una siguiente iteracion.",
            "No se han incorporado credenciales ni notas internas del documento antiguo, por seguridad y limpieza academica.",
        ],
    )

    document.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_document()
